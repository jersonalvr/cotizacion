# app.py
import streamlit as st
import requests
from rembg import remove
from PIL import Image
import os
from geopy.geocoders import Nominatim
from streamlit_js_eval import get_geolocation
import folium
from streamlit_folium import st_folium
import pdfplumber
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from datetime import datetime
from io import BytesIO
import zipfile
import pyperclip
from st_copy_to_clipboard import st_copy_to_clipboard
from streamlit_image_comparison import image_comparison

# Determinar la ruta base de la aplicación
base_dir = os.path.dirname(os.path.abspath(__file__))

def obtener_datos_sunat(dni):
    apisnet_key = st.secrets["APISNET"]["key"]
    url = f"https://api.apis.net.pe/v2/sunat/dni?numero={dni}&token={apisnet_key}"
    try:
        response = requests.get(url)
        if response.status_code == 200:
            data = response.json()
            nombres = f"{data.get('nombres', '')} {data.get('apellidoPaterno', '')} {data.get('apellidoMaterno', '')}".strip()
            ruc = data.get("ruc", "")
            return nombres, ruc
        else:
            st.error("Error al obtener datos de SUNAT. Verifica el DNI ingresado.")
            return None, None
    except Exception as e:
        st.error(f"Error al conectar con la API de SUNAT: {e}")
        return None, None

def obtener_direccion_desde_coordenadas(lat, lon):
    geolocator = Nominatim(user_agent="my_streamlit_app")
    try:
        location = geolocator.reverse((lat, lon))
        return location.address
    except Exception as e:
        st.error(f"Error al obtener la dirección: {e}")
        return None

def crear_mapa(lat=None, lon=None, zoom=13):
    # Usar coordenadas proporcionadas o predeterminadas de Lima, Perú
    if lat is None or lon is None:
        lat, lon = -12.0464, -77.0428  # Coordenadas de Lima

    # Crear el mapa centrado en la ubicación y usando el zoom proporcionado
    m = folium.Map(location=[lat, lon], zoom_start=zoom)

    # Agregar marcador si hay coordenadas específicas
    folium.Marker(
        [lat, lon],
        popup="Ubicación actual",
        icon=folium.Icon(color='red', icon='info-sign'),
        draggable=False
    ).add_to(m)

    return m

def extraer_nombre_servicio(pdf_file):
    texto_completo = ""
    with pdfplumber.open(pdf_file) as pdf:
        for pagina in pdf.pages:
            texto_completo += pagina.extract_text()

    texto_unido = ' '.join(texto_completo.split())

    patron = r'2\.\s*OBJETO\s*DE\s*LA\s*CONTRATACION\s*(.*?)\s*3\.\s*FINALIDAD\s*PUBLICA'

    match = re.search(patron, texto_unido, re.DOTALL | re.IGNORECASE)

    if match:
        servicio = ' '.join(match.group(1).split())
        return servicio
    return "Servicio no encontrado"

def extraer_forma_pago(pdf_file):
    texto_completo = ""
    with pdfplumber.open(pdf_file) as pdf:
        for pagina in pdf.pages:
            texto_completo += pagina.extract_text()

    texto_unido = ' '.join(texto_completo.split())

    patron = r'El pago se realizará en\s*(.*?)\s*luego de la emisión de la conformidad del servicio,'

    match = re.search(patron, texto_unido, re.DOTALL | re.IGNORECASE)

    if match:
        forma_pago = ' '.join(match.group(1).split()).upper()
        return forma_pago
    return "FORMA DE PAGO NO ENCONTRADA"

def extraer_dias(pdf_file):
    texto_completo = ""
    with pdfplumber.open(pdf_file) as pdf:
        for pagina in pdf.pages:
            texto_completo += pagina.extract_text()

    texto_unido = ' '.join(texto_completo.split())

    patron = r'El plazo de ejecución del servicio es de hasta\s*(\d+)\s*días calendario'

    match = re.search(patron, texto_unido, re.DOTALL | re.IGNORECASE)

    if match:
        dias = match.group(1)
        return dias
    return "DÍAS NO ENCONTRADOS"

def procesar_firma(firma_file, remover_fondo=False):
    """
    Procesa la imagen de la firma, opcionalmente removiendo el fondo.
    
    Args:
        firma_file: Archivo de imagen subido
        remover_fondo: Boolean indicando si se debe remover el fondo
    
    Returns:
        BytesIO: Imagen procesada en formato BytesIO
    """
    # Abrir la imagen
    image = Image.open(firma_file)
    
    if remover_fondo:
        with st.spinner('Removiendo fondo de la firma...'):
            # Remover fondo
            imagen_procesada = remove(image)
            # Convertir a modo RGBA si no lo está ya
            if imagen_procesada.mode != 'RGBA':
                imagen_procesada = imagen_procesada.convert('RGBA')
    else:
        imagen_procesada = image
        
    # Convertir a BytesIO
    img_byte_arr = BytesIO()
    imagen_procesada.save(img_byte_arr, format='PNG')
    img_byte_arr.seek(0)
    
    return img_byte_arr

def mostrar_seccion_firma():
    """
    Muestra la sección de carga y procesamiento de firma
    
    Returns:
        BytesIO: Imagen de firma procesada
        bool: Indicador de si se cargó una firma
    """
    st.header("Sube tu firma")
    
    # Checkbox para remover fondo
    remover_fondo = st.checkbox("Remover fondo de la firma", value=False)
    
    # Upload de firma
    firma_file = st.file_uploader(
        "Selecciona tu imagen de firma", 
        type=["png", "jpg", "jpeg"],
        help="Sube una imagen de tu firma en formato PNG, JPG o JPEG"
    )
    
    if firma_file is not None:
        # Procesar firma
        firma_procesada = procesar_firma(firma_file, remover_fondo)
        
        if remover_fondo:
            # Mostrar comparación antes/después
            col1, col2 = st.columns(2)
            with col1:
                st.write("Firma Original")
                st.image(firma_file, width=300)
            with col2:
                st.write("Firma sin fondo")
                st.image(firma_procesada, width=300)
                
            # Opcionalmente mostrar comparador deslizante
            st.write("Comparador deslizante")
            image_comparison(
                img1=Image.open(firma_file),
                img2=Image.open(firma_procesada),
                label1="Original",
                label2="Sin fondo"
            )
        else:
            # Mostrar solo la firma original
            st.image(firma_file, caption="Vista previa de la firma", width=300)
        
        return firma_procesada, True
    
    return None, False

def generar_cotizacion(pdf_file, data):
    # Extraer datos del PDF
    servicio = extraer_nombre_servicio(pdf_file)
    forma_pago = extraer_forma_pago(pdf_file)
    dias = extraer_dias(pdf_file)

    # Actualizar data con los datos extraídos
    data['servicio'] = servicio
    data['armada'] = forma_pago
    data['dias'] = dias

    # Cargar el documento
    template_path = os.path.join(base_dir, 'FormatoCotizacion.docx')
    doc = Document(template_path)

    # Diccionario de reemplazos
    reemplazos = {
        '{{fecha}}': data['fecha'],
        '{{servicio}}': data['servicio'],
        '{{dias}}': data['dias'],
        '{{oferta}}': "{:.2f}".format(data['oferta']),
        '{{armada}}': data['armada'],
        '{{MES}}': data['mes'],
        '{{dni}}': data['dni'],
        '{{nombres}}': data['nombres'],
        '{{ruc}}': data['ruc'],
        '{{telefono}}': data['telefono'],
        '{{correo}}': data['correo'],
        '{{direccion}}': data['direccion'],
        '{{banco}}': data['banco'],
        '{{cuenta}}': data['cuenta'],
        '{{cci}}': data['cci'],
        '{{year}}': str(data['year']),
    }

    def reemplazar_texto(texto, reemplazos):
        for key, value in reemplazos.items():
            texto = texto.replace(key, str(value))
        return texto

    def procesar_parrafo(paragraph):
        if '{{firma}}' in paragraph.text:
            # Manejar la firma como antes
            p = paragraph._element
            p.clear_content()
            run = paragraph.add_run()
            data['firma'].seek(0)
            run.add_picture(BytesIO(data['firma'].read()), height=Cm(1.91))
        else:
            # Concatenar todo el texto de los runs en el párrafo
            full_text = ''
            formatting = []
            for run in paragraph.runs:
                full_text += run.text
                formatting.append({
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    # No guardamos font_name ni font_size
                    'font_color': run.font.color.rgb
                })

            # Reemplazar los marcadores de posición en el texto completo
            new_full_text = reemplazar_texto(full_text, reemplazos)

            # Borrar los runs existentes
            for run in paragraph.runs:
                run.text = ''

            # Crear un nuevo run con el texto reemplazado
            run = paragraph.add_run(new_full_text)
            # Aplicar el formato del primer run original
            if formatting:
                fmt = formatting[0]
                run.bold = fmt['bold']
                run.italic = fmt['italic']
                run.underline = fmt['underline']
                run.font.color.rgb = fmt['font_color']
            else:
                # Valores por defecto si no hay formato original
                run.bold = False
                run.italic = False
                run.underline = False

            # Establecer la fuente a Arial 11
            run.font.name = 'Arial'
            run.font.size = Pt(11)

    # Procesar todos los párrafos en el documento principal
    for paragraph in doc.paragraphs:
        procesar_parrafo(paragraph)

    # Función recursiva para procesar tablas anidadas
    def procesar_tabla(tabla):
        for row in tabla.rows:
            for cell in row.cells:
                # Procesar párrafos dentro de la celda
                for paragraph in cell.paragraphs:
                    procesar_parrafo(paragraph)

                # Procesar tablas anidadas dentro de la celda
                for tabla_anidada in cell.tables:
                    procesar_tabla(tabla_anidada)

    # Procesar todas las tablas en el documento
    for tabla in doc.tables:
        procesar_tabla(tabla)

    # Guardar el documento modificado en un BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

def generar_cci(banco, cuenta):
    if not banco or not cuenta or banco == "Otros":
        return ""
    
    cuenta_limpia = cuenta.replace("-", "")
    cci_map = {
        "BCP": "002" + cuenta_limpia + "13",
        "Interbank": "003" + cuenta_limpia + "43",
        "Scotiabank": "00936020" + cuenta_limpia + "95",
        "Banco de la Nación": "0187810" + cuenta_limpia + "55",
        "BanBif": "0386501" + cuenta_limpia + "83"
    }
    return cci_map.get(banco, "")

def crear_donation_footer(base_dir):
    footer = st.container()
    
    with footer:
        st.markdown("---")
        st.header("💝 Apoya este proyecto")
        
        # Tabs for different payment methods
        tab1, tab2, tab3, tab4, tab5 = st.tabs(["Yape", "Depósito Bancario", "Tarjeta", "Otros Métodos", "Crypto"])
        
        # Yape Tab
        with tab1:
            st.subheader("Donar por Yape")
            col1, col2 = st.columns([1, 1])
            with col1:
                yape_image_path = os.path.join(base_dir, "yape.png")
                if os.path.exists(yape_image_path):
                    st.image(yape_image_path, width=300)
                else:
                    st.error(f"No se encontró la imagen en: {yape_image_path}")
            with col2:
                # Agregamos el número de Yape con botón para copiar
                st_copy_to_clipboard("964536063", "Copiar número de Yape")

        # Bank Deposits Tab
        with tab2:
            col1, col2 = st.columns([1, 1])
            
            with col1:
                st.subheader("En Soles 🇵🇪")
                banks_soles = {
                    "BCP": "31004315283063",
                    "BanBif": "008023869670",
                    "Interbank": "8983223094904",
                    "Scotiabank": "7508188435"
                }
                
                for bank, account in banks_soles.items():
                    st.write(f"**{bank}:**")
                    st_copy_to_clipboard(account, f"Copiar cuenta {bank}")
            
            with col2:
                st.subheader("En Dólares 💵")
                banks_usd = {
                    "BCP": "31004319160179",
                    "Interbank": "8983224537574",
                    "Scotiabank": "7508188145"
                }
                
                for bank, account in banks_usd.items():
                    st.write(f"**{bank}:**")
                    st_copy_to_clipboard(account, f"Copiar cuenta {bank}")
        
        # Card Payments Tab
        with tab3:
            st.subheader("Donar con tarjeta 💳")
            
            amounts = {
                "10": "https://pago-seguro.vendemas.com.pe/MTYzNDc3OTY3NjYxMWM4MjU2MTIuNzMxNzMxMjgxNTUz",
                "15": "https://pago-seguro.vendemas.com.pe/ZjE5MTc2MWIzMjM0MDQ3NDQ0NC4yOWQxNzMxMjgxNjIz",
                "20": "https://pago-seguro.vendemas.com.pe/NmEzOTkwNzI1OTY1Zi42MTM0MDg2NDYxNzMxMjgxNjUy",
                "25": "https://pago-seguro.vendemas.com.pe/MTUzODVjYTM4NDIzZDgxNjMwLjI0NzcxNzMxMjgxNjc3",
                "30": "https://pago-seguro.vendemas.com.pe/ODM0LjIyMTYyMjA2NjZmYjdhNmMzM2QxNzMxMjgxNzA3",
                "35": "https://pago-seguro.vendemas.com.pe/ODM3MzMzMDFkNDcuOTE2YzUzMjQxNTExNzMxMjgxNzI1",
                "40": "https://pago-seguro.vendemas.com.pe/YzgwZjM2NzM1LjZiMWQ0NzEzNTMxNzYxNzMxMjgxNzQ0",
                "45": "https://pago-seguro.vendemas.com.pe/MTgwMTYuNzQ1MzM0OTk0MzI4MTQ2MzYxNzMxMjgxNzYy"
            }
            
            # Creamos dos filas de 4 columnas cada una para mejor visualización
            for row in range(2):
                cols = st.columns(4)
                start_idx = row * 4
                end_idx = start_idx + 4
                
                # Tomamos solo los montos correspondientes a esta fila
                row_amounts = dict(list(amounts.items())[start_idx:end_idx])
                
                for col_idx, (amount, link) in enumerate(row_amounts.items()):
                    with cols[col_idx]:
                        st.link_button(f"S/ {amount}", link)
            
            st.link_button("Más de S/ 50", "https://linkdecobro.ligo.live/v3/44df73097f594239b21b78b6905bed98")
        
        # Other Payment Methods Tab
        with tab4:
            col1, col2 = st.columns([1, 1])
            
            with col1:
                st.subheader("Mercado Pago")
                st.link_button("Donar con Mercado Pago", "https://link.mercadopago.com.pe/jersonapp")
            
            with col2:
                st.subheader("PayPal")
                st.link_button("Donar con PayPal", "https://www.paypal.com/paypalme/dschimbote")
        
        # Crypto Tab
        with tab5:
            st.subheader("Binance")
            col1, col2 = st.columns([1, 1])
            
            with col1:
                binance_image_path = os.path.join(base_dir, "binance.png")
                if os.path.exists(binance_image_path):
                    st.image(binance_image_path, width=300)
                else:
                    st.error(f"No se encontró la imagen en: {binance_image_path}")
            
            with col2:
                st.link_button("Donar con Binance", "https://app.binance.com/qr/dplkbb7f88c5329c4692adf278670d1b37ab")
                
def main():
    st.set_page_config(
        page_title="Genera tu Cotización",
        page_icon="🎣",
        layout="wide"
    )
    # Inicializar variables de estado para la ubicación
    if 'zoom' not in st.session_state:
        st.session_state['zoom'] = 13
    if 'lat' not in st.session_state:
        st.session_state['lat'] = None
    if 'lon' not in st.session_state:
        st.session_state['lon'] = None
    if 'direccion' not in st.session_state:
        st.session_state['direccion'] = ''

    # Sección de carga de TDR
    st.header("Sube tu TDR (PDF)")
    pdf_file = st.file_uploader("Selecciona tu archivo PDF", type=["pdf"])

    # Sección de firma
    firma_procesada, firma_cargada = mostrar_seccion_firma()

    # Inicializar variables de estado si no existen
    if 'form_data' not in st.session_state:
        st.session_state.form_data = {
            'dni': '',
            'nombres': '',
            'ruc': '',
            'telefono': '',
            'correo': '',
            'direccion': '',
            'banco': '',
            'cuenta': '',
            'cci': '',
            'oferta': 0.0
        }

    # Sección de datos personales
    st.header("Datos Personales")

    # DNI y datos de SUNAT
    dni = st.text_input("Introduce tu DNI", max_chars=8, key='dni_input')
    if dni and len(dni) == 8:
        nombres, ruc = obtener_datos_sunat(dni)
        if nombres:
            st.success(f"Nombres: {nombres}")
            st.success(f"RUC: {ruc}")
            # Guardar en session_state
            st.session_state.form_data.update({
                'dni': dni,
                'nombres': nombres,
                'ruc': ruc
            })

    # Información de contacto
    col1, col2 = st.columns(2)

    with col1:
        telefono = st.text_input("Teléfono", key='telefono_input')
        if telefono:
            st.session_state.form_data['telefono'] = telefono

    with col2:
        correo = st.text_input("Correo electrónico", key='correo_input')
        if correo:
            st.session_state.form_data['correo'] = correo

    st.subheader("Dirección")
    col1, col2 = st.columns([1, 1])

    with col1:
        if st.button("Obtener ubicación actual"):
            get_geolocation('geo_loc')
            st.write("Obteniendo ubicación...")

    # Recuperar la ubicación desde st.session_state después de la llamada
    if 'geo_loc' in st.session_state:
        loc = st.session_state['geo_loc']
        if loc and 'coords' in loc:
            st.session_state.lat = loc['coords']['latitude']
            st.session_state.lon = loc['coords']['longitude']
            direccion = obtener_direccion_desde_coordenadas(
                st.session_state.lat,
                st.session_state.lon
            )
            if direccion:
                st.session_state.direccion = direccion

    # Un solo campo de dirección fuera de las columnas
    direccion_input = st.text_input(
        "Dirección",
        value=st.session_state.get('direccion', ''),
        key="direccion_input"
    )
    # Actualizar el estado con el valor del input
    st.session_state.direccion = direccion_input

    with col2:
        # Mostrar el mapa con la ubicación si está disponible y el zoom actual
        mapa = crear_mapa(
            lat=st.session_state['lat'],
            lon=st.session_state['lon'],
            zoom=st.session_state['zoom']
        )
        mapa_data = st_folium(
            mapa,
            height=300,
            width=None,
            returned_objects=["last_clicked", "zoom"]
        )

        # Actualizar ubicación cuando se hace clic en el mapa
        if mapa_data["last_clicked"]:
            clicked_lat = mapa_data["last_clicked"]["lat"]
            clicked_lng = mapa_data["last_clicked"]["lng"]

            # Guardar el zoom actual antes de actualizar
            if mapa_data.get("zoom"):
                st.session_state['zoom'] = mapa_data["zoom"]

            # Actualizar estado
            st.session_state['lat'] = clicked_lat
            st.session_state['lon'] = clicked_lng
            nueva_direccion = obtener_direccion_desde_coordenadas(clicked_lat, clicked_lng)
            if nueva_direccion:
                st.session_state['direccion'] = nueva_direccion
                st.rerun()
        # Actualizar el zoom incluso si no se hace clic
        elif mapa_data.get("zoom"):
            st.session_state['zoom'] = mapa_data["zoom"]

    # Información bancaria
    st.header("Información Bancaria")
    banco_seleccionado = st.selectbox(
        "Selecciona tu banco",
        ["BCP", "Interbank", "Scotiabank", "Banco de la Nación", "BanBif", "Otros"],
        key='banco_input'
    )
    if banco_seleccionado:
        st.session_state.form_data['banco'] = banco_seleccionado

    cuenta = st.text_input("Ingresa tu cuenta", max_chars=20, key='cuenta_input')
    if cuenta:
        st.session_state.form_data['cuenta'] = cuenta

    # Generar y mostrar CCI
    cci = st.text_input("CCI (editable)", value=generar_cci(banco_seleccionado, cuenta), key='cci_input')
    if cci:
        st.session_state.form_data['cci'] = cci
        
    # Sección de oferta económica
    st.header("Oferta Económica")
    
    # Extraer días del PDF si está disponible
    dias = "30"  # Valor por defecto
    if pdf_file:
        dias = extraer_dias(pdf_file)
    
    # Obtener el valor sugerido basado en los días
    try:
        dias_int = int(dias)
        valor_sugerido = ((dias_int - 1) // 30 + 1) * 2000.00
    except ValueError:
        valor_sugerido = 2000.00
    
    col1, col2 = st.columns([3, 1])
    with col1:
        oferta_total = st.number_input(
            "OFERTA TOTAL (S/)",
            min_value=0.0,
            value=valor_sugerido,  # Valor sugerido dinámico
            step=10.0,
            format="%.2f",
            help=f"Valor sugerido: S/ {valor_sugerido:,.2f} para {dias} días. Puedes ajustar el monto usando las flechas (±10) o ingresando directamente el valor deseado."
        )
    
    with col2:
        st.markdown("""
        <style>
        .small-font {
            font-size: 0.9em;
            color: #666;
        }
        </style>
        """, unsafe_allow_html=True)
        
        st.markdown(
            f'<p class="small-font">Sugerido: S/ {valor_sugerido:,.2f}<br>Incrementos: ± S/ 10.00</p>', 
            unsafe_allow_html=True
        )
    
    # Mostrar el valor ingresado con formato de moneda
    if oferta_total > 0:
        st.write(f"Monto ingresado: S/ {oferta_total:,.2f}")

    # Botón de envío
    if st.button("Generar cotizacion"):
        if not all([pdf_file, firma_cargada, dni, st.session_state.direccion, telefono, correo, banco_seleccionado, cuenta, cci, oferta_total]):
            st.error("Por favor, complete todos los campos requeridos.")
        else:
            # Obtener datos de SUNAT
            nombres, ruc = obtener_datos_sunat(dni)
            if not nombres:
                st.error("No se pudo obtener datos de SUNAT. Verifica el DNI ingresado.")
            else:
                # Formatear la fecha actual en español
                meses = {
                    "January": "enero", "February": "febrero", "March": "marzo", "April": "abril",
                    "May": "mayo", "June": "junio", "July": "julio", "August": "agosto",
                    "September": "setiembre", "October": "octubre", "November": "noviembre", "December": "diciembre"
                }
                fecha_actual = datetime.now()
                mes_actual = meses[fecha_actual.strftime("%B")].upper()
                fecha_formateada = f"{fecha_actual.day} de {meses[fecha_actual.strftime('%B')]} de {fecha_actual.year}"

                # Preparar datos para generar la cotización
                data = {
                    'dni': dni,
                    'nombres': nombres,
                    'ruc': ruc,
                    'telefono': telefono,
                    'correo': correo,
                    'direccion': st.session_state.direccion,
                    'banco': banco_seleccionado,
                    'cuenta': cuenta,
                    'cci': cci,
                    'oferta': oferta_total,
                    'fecha': fecha_formateada,
                    'year': fecha_actual.year,
                    'mes': mes_actual,
                    'firma': firma_procesada,  # Usar la firma procesada
                }

                # Generar la cotización
                doc_io = generar_cotizacion(pdf_file, data)

                # Crear un archivo ZIP en memoria
                zip_io = BytesIO()
                with zipfile.ZipFile(zip_io, mode='w', compression=zipfile.ZIP_DEFLATED) as zipf:
                    # Agregar el documento de cotización
                    zipf.writestr('Formato de Cotización.docx', doc_io.getvalue())
                    # Agregar la firma
                    firma_procesada.seek(0)  # Reiniciar el puntero del archivo
                    zipf.writestr('Firma.png', firma_procesada.getvalue())
                    # Agregar el TDR original
                    pdf_file.seek(0)  # Reiniciar el puntero del archivo
                    zipf.writestr('6. Copia de Terminos de Referencia.pdf', pdf_file.getvalue())

                zip_io.seek(0)
                st.success("¡Cotización generada correctamente!")

                # Botón para descargar el ZIP
                st.download_button(
                    label="Descargar Todos los Archivos Generados (ZIP)",
                    data=zip_io.getvalue(),
                    file_name="cotizacion.zip",
                    mime="application/zip",
                )
                
    st.markdown("""
        <h3 style='text-align: center; margin-bottom: 2rem;'>
            Descarga ahora el generador de constancias RNP, RUC, RNSSC
        </h3>
    """, unsafe_allow_html=True)

    # Crear dos columnas
    col1, col2 = st.columns(2)

    # Columna izquierda - Instrucciones
    with col1:
        st.markdown("### 📝 Instrucciones de uso:")
        st.markdown("""
        1. **Requisitos previos:**
        - Google Chrome instalado en su equipo
        - Conexión a internet estable
        
        2. **Proceso de instalación:**
        - Si no tiene Chrome, descárguelo primero
        - Descargue el Generador de Constancias
        - Ejecute el archivo descargado
        
        3. **Funcionalidades:**
        - Genera constancias RNP
        - Obtiene información RUC
        - Consulta RNSSC
        """)

    # Columna derecha - Botones y advertencias
    with col2:
        st.warning("⚠️ Requiere Google Chrome para funcionar")
        
        st.link_button(
            "🌐 Descargar Google Chrome",
            "https://www.google.com/intl/es-419/chrome/dr/download/",
            use_container_width=True,
        )
        
        st.markdown("<br>", unsafe_allow_html=True)
        
        st.info("🔒 Aplicación verificada y segura")
        
        file_url = "https://drive.usercontent.google.com/download?id=1084eOd4CSqMQ323U1-walYGELyvo6yei&export=download&confirm=t&uuid=5acc3199-ccbb-4fe3-86fd-62de9bddfca7"
        st.download_button(
            label="📥 Descargar Generador de Constancias",
            data=requests.get(file_url).content,
            file_name="constancia.exe",
            mime="application/x-msdownload",
            use_container_width=True,
        )
        
    crear_donation_footer(base_dir)
    
if __name__ == "__main__":
    main()
